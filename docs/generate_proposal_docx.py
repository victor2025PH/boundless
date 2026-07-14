# -*- coding: utf-8 -*-
"""
无界科技 BOUNDLESS —— 全球全自动社媒引流与 AI 转化系统 · 部署方案书 (Word)
与官网方案页 /proposal/ 内容对齐。使用 python-docx 生成。
"""
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 单一事实源：费用/配置/产能/平台统一从 proposal_data 取，避免与网页数字漂移。
sys.path.insert(0, str(Path(__file__).resolve().parent))
import proposal_data as D  # noqa: E402

OUT = Path(__file__).with_name("无界全球全自动引流系统_部署方案书.docx")
DEPLOY = D.FEES["deploy"]   # "11,999"
MONTH = D.FEES["month"]     # "3,000"

# ── 品牌配色 ──
INDIGO = "3730A3"     # 主标题（靛蓝紫）
CYAN = "0891B2"       # 强调青
DARK = "1E1B4B"       # 深色正文标题
VIOLET = "6D28D9"
GRAY = "F2F3F9"
L_CYAN = "E0F2FE"
L_GREEN = "E3F5E9"
L_YELLOW = "FEF3D7"
L_VIOLET = "EDE9FE"
FONT = "Microsoft YaHei"


def _set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    _set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_fill=INDIGO, right_cols=None):
    right_cols = right_cols or []
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        al = WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else None
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF", size=9.5, align=al)
        set_cell_shading(hdr[i], header_fill)
        if widths:
            hdr[i].width = Cm(widths[i])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        is_total = isinstance(row, dict)
        data = row["cells"] if is_total else row
        for i, val in enumerate(data):
            al = WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else None
            set_cell_text(cells[i], val, bold=is_total,
                          color=(CYAN if is_total else None), align=al)
            if is_total:
                set_cell_shading(cells[i], L_CYAN)
            elif r_idx % 2 == 1:
                set_cell_shading(cells[i], "F8F9FE")
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_font(run, color=(INDIGO if level <= 2 else DARK))
    return p


def add_para(doc, text, size=10.5, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    _set_font(run, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_font(run, size=10)


def add_callout(doc, title, body, fill=L_VIOLET):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _set_font(r, size=11, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    _set_font(r2, size=10, color="333333")
    doc.add_paragraph()


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = FONT
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # ── 封面 ──
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.space_before = Pt(90)
    r = logo.add_run("无界科技  BOUNDLESS")
    _set_font(r, size=16, bold=True, color=VIOLET)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("让沟通，无界")
    _set_font(r, size=11, color="888888")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    r = title.add_run("全球全自动社媒引流\n与 AI 转化系统")
    _set_font(r, size=30, bold=True, color=INDIGO)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(14)
    r = subtitle.add_run("部 署 与 服 务 方 案 书")
    _set_font(r, size=15, color=CYAN)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(30)
    r = line.add_run("真机集群 RPA  ×  多语言 AI 客服  ×  引流-承接-转化一体化闭环")
    _set_font(r, size=11, color="555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(150)
    r = meta.add_run("文档性质：客户专属商业方案（保密）\n"
                     "覆盖平台：TikTok / Facebook / Messenger（已成熟）\n"
                     "版本：V1.0     日期：2026-07")
    _set_font(r, size=11, color="444444")

    doc.add_section(WD_SECTION.NEW_PAGE)

    # ── 一、方案概述 ──
    add_heading(doc, "一、方案概述", 1)
    add_callout(
        doc, "一句话定位",
        "以真实手机集群 RPA 驱动 TikTok / Facebook 等全球社交平台，7×24 小时全自动养号、"
        "获客、触达；再由多语言 AI 客服无缝承接对话、推进成交。一套系统打通「引流 → 承接 → 转化」全链路，"
        "让 1 名运营即可管理数十至上百个海外账号。", L_GREEN)
    add_para(doc, "本方案面向出海企业的规模化获客场景，交付的不是单纯的「群控工具」，而是从流量触达到成交"
                  "承接的完整增长系统。核心引流平台 TikTok 与 Facebook/Messenger 已成熟稳定投产，其余主流平台"
                  "（Instagram / WhatsApp / Telegram / Twitter·X / LinkedIn / LINE）已具备同栈基础能力，"
                  "如需作为正式引流渠道，可按需二次开发快速接入。")
    add_table(doc, ["核心维度", "说明"], [
        ["产品定位", "获客执行 + AI 转化承接一体化增长系统"],
        ["已成熟平台", "TikTok、Facebook、Messenger（全链路投产）"],
        ["可扩展平台", "Instagram / WhatsApp / Telegram / X / LinkedIn / LINE（二次开发）"],
        ["技术路线", "真机 RPA（ADB + UIAutomator2），不走平台 API，风控识别率低"],
        ["AI 能力", "四层触发 + 多维情绪识别 + 意向评分，15+ 语言拟人应答"],
        ["部署方式", "私有化部署，数据完全隔离，主控 + 多 Worker 分布式扩展"],
        ["部署搭建费", f"{DEPLOY} {D.FEES['currency']}（一次性）"],
        ["月度服务费", f"{MONTH} {D.FEES['currency']} / 月"],
        ["后续开发", "超出现有功能范围的新需求，独立评估报价"],
    ], widths=[4.5, 11.5])

    # ── 二、系统能力与技术架构 ──
    add_heading(doc, "二、系统能力与技术架构", 1)
    add_para(doc, "系统由「获客执行层」与「AI 转化承接层」两大自研核心组成，通过统一数据与事件总线衔接，"
                  "形成从触达到成交的闭环。")
    add_table(doc, ["能力模块", "关键特性"], [
        ["真机 RPA 引擎", "ADB + UIAutomator2 纯 UI 模拟；VLM 视觉兜底；不走 API / 网页协议"],
        ["AI 决策与应答", "四层触发、情绪多维识别（信任/兴趣/焦虑）、意向评分、知识库直答"],
        ["反风控合规栈", "拟人行为曲线、配额限流、冷却、VPN/代理伪装、设备指纹隔离"],
        ["集群与设备", "主控 + N Worker 分布式；热插拔纳管；Scrcpy 实时投屏群控（<20ms）"],
        ["数据与看板", "转化漏斗、客户画像、SLA 看板、A/B 实验自动选优、周报/日报"],
        ["监控与告警", "设备在线率、风控事件熔断、Telegram/飞书/钉钉告警、成本追踪"],
    ], widths=[4.5, 11.5])
    add_callout(doc, "扩展性说明",
                "单台工作机（Worker）通过独立供电的工业 USB HUB 可稳定管理约 15–16 台真机；"
                "通过增加 Worker 节点横向扩展，单个主控集群可管理 200+ 台设备。", L_CYAN)

    # ── 三、全自动引流闭环 ──
    add_heading(doc, "三、全自动引流闭环（六步）", 1)
    add_table(doc, ["步骤", "环节", "系统动作"], [
        ["1", "智能养号", "模拟真人浏览点赞，建立兴趣标签与账号权重（约 7 天冷启动）"],
        ["2", "精准获客", "关键词 / 话题标签 / 竞品粉丝 / 评论区 / 群成员采集目标客户"],
        ["3", "自动触达", "加好友 / 关注 / 互动 / 私信，个性化话术批量执行"],
        ["4", "AI 承接", "多语言智能客服自动应答，识别意向、反异议、FAQ 直答"],
        ["5", "引流转化", "引导至 WhatsApp / Telegram / 私域，高意向自动转人工接管"],
        ["6", "数据复盘", "漏斗 / 成本 / 转化全程可视，A/B 实验持续优化"],
    ], widths=[1.6, 3.4, 11], right_cols=[0])

    # ── 四、平台能力矩阵 ──
    add_heading(doc, "四、平台能力矩阵", 1)
    add_para(doc, "核心引流平台已全链路打通并稳定投产；其余平台已具备同栈基础能力，作为正式引流渠道需二次开发。")
    add_table(doc, ["平台", "主要能力", "状态"],
              [[name, cap, D.STATUS_LABEL[st]] for name, cap, st in D.PLATFORMS],
              widths=[3.5, 9.5, 3])

    # ── 五、核心功能清单 ──
    add_heading(doc, "五、核心功能清单", 1)
    add_table(doc, ["功能模块", "包含能力"], [
        ["获客引擎", "关键词/话题搜索、竞品粉丝采集、评论区截流、群成员提取、Lead 库去重"],
        ["AI 智能转化", "四层触发+意向评分、多维情绪识别、15+ 语言拟人回复、FAQ 直答、转人工"],
        ["反风控与合规", "拟人行为曲线、分平台配额限流、代理池+出口 IP 检查、指纹隔离、渐进放量"],
        ["集群设备管理", "主控+N Worker、热插拔纳管、Scrcpy 群控、掉线自动恢复、批量 APK/账号分配"],
        ["数据运营看板", "转化漏斗、客户画像与旅程、SLA/KPI、A/B 自动选优、周报日报"],
        ["监控告警", "在线率/心跳、风控熔断、多渠道告警、成本追踪、E2E 自检"],
    ], widths=[4, 12])

    doc.add_section(WD_SECTION.NEW_PAGE)

    # ── 六、硬件配置方案 ──
    add_heading(doc, "六、硬件配置方案", 1)
    add_callout(doc, "重要说明",
                "以下为推荐硬件配置清单，属客户一次性资产投入，独立于软件部署费与月服务费；可自行采购，"
                "亦可委托我方代采整备。硬件价格每日波动，本方案仅列配置与数量、不标注单价与总价，"
                "请以采购当日市场行情为准。", L_YELLOW)

    for idx, tier in enumerate(D.TIERS):
        letter = "ABC"[idx] if idx < 3 else str(idx + 1)
        add_heading(doc, f"配置 {letter} · {tier['label']}"
                         f"（{tier['phones']} 台真机 · {tier['topo']}）", 2)
        add_table(doc, ["类别", "型号 / 规格", "数量"],
                  [list(row) for row in tier["bom"]],
                  widths=[3.5, 10.5, 2], right_cols=[2])

    add_callout(doc, "选型说明",
                "① 手机推荐 Redmi 13C 等中端安卓机（性价比高、UI 稳定、适配成熟），亦可换用三星 A 系列；"
                "② 主控是否配备 RTX 4090 取决于是否启用 AI 实时换脸 / 本地私有大模型，仅做引流可选标准 GPU 降低成本；"
                "③ 建议使用稳定 Wi-Fi + 目标地区住宅/移动代理 IP，避免依赖单一网络。", L_CYAN)

    # ── 七、配置与产能测算 ──
    add_heading(doc, "七、配置与产能测算", 1)
    add_para(doc, "下表为不同规模的成熟期理论产能估算（单机日均：主动触达约 650、私信约 280、采集线索约 1000；"
                  "月意向对话按私信触达的约 15% 估算）。用于规模规划参考，不构成业绩承诺。")
    def _prow(label, phones, machines):
        p = D.productivity(phones)
        return [label, str(phones), str(machines),
                D.fmt(p["reach"]), D.fmt(p["dm"]), D.fmt(p["scrape"]), D.fmt(p["leads"])]

    import math as _math
    _prod_rows = [_prow(t["label"].split()[0], t["phones"], t["machines"]) for t in D.TIERS]
    _prod_rows.append(_prow("超大规模", D.XLARGE_PHONES,
                            _math.ceil(D.XLARGE_PHONES / D.PER_MACHINE)))
    add_table(doc, ["规模", "真机数", "主机数", "月主动触达", "月私信触达", "月采集线索", "月意向对话"],
              _prod_rows,
              widths=[2.6, 1.9, 1.9, 2.5, 2.5, 2.5, 2.4],
              right_cols=[1, 2, 3, 4, 5, 6])

    doc.add_section(WD_SECTION.NEW_PAGE)

    # ── 八、费用方案 ──
    add_heading(doc, "八、费用方案", 1)
    add_callout(doc, "三段式报价",
                "软件部署与服务费用独立于硬件资产。当前已成熟功能范围内一价交付，后续新增需求另行评估报价。",
                L_VIOLET)

    add_heading(doc, f"① 部署搭建费：{DEPLOY} {D.FEES['currency']}（一次性）", 2)
    for t in [
        "系统私有化部署（主控 + 全部 Worker 节点）",
        "TikTok / Facebook / Messenger 引流链路配置",
        "设备集群接入、账号纳管、投屏群控调通",
        "AI 客服话术 / 知识库 / 多语言初始化",
        "反风控参数、配额、代理 / VPN 配置",
        "数据看板与告警接入",
        "端到端联调 + 1 对 1 运营培训",
    ]:
        add_bullet(doc, t)

    add_heading(doc, f"② 月度服务费：{MONTH} {D.FEES['currency']} / 月", 2)
    for t in [
        "系统运维与稳定性保障",
        "平台 UI 更新的选择器 / 适配跟进",
        "反风控策略持续调优",
        "版本更新与现有功能范围内的小迭代",
        "话术 / 知识库优化建议",
        "优先技术支持与故障响应",
        "每周数据 review（可选）",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "③ 二次开发费：按需独立报价", 2)
    for t in [
        "新增平台接入（Instagram / WhatsApp / X / LinkedIn / LINE 等）",
        "定制业务工作流 / 行业域包",
        "AI 实时换脸 / 声音克隆 / 数字人",
        "深度 CRM / 广告归因 / BI 定制",
        "多租户 / 私有云 SaaS 化改造",
        "合规能力扩展（GDPR / CCPA）",
    ]:
        add_bullet(doc, t)

    add_callout(doc, "费用边界",
                "「部署费 + 月服务费」覆盖当前已成熟的 TikTok / Facebook / Messenger 全自动引流与 AI 转化能力。"
                "硬件（主机 / 手机 / 配件）为客户一次性资产，可自备或委托代采。任何超出现有功能的新需求，"
                "我方将先出具需求评估与独立开发报价，双方确认后再启动。", L_YELLOW)

    # ── 九、ROI 测算示例 ──
    add_heading(doc, "九、ROI 投资回报测算示例", 1)
    add_para(doc, "以专业版（60 台真机）为例，假设客单价 300 USD、毛利率 50%、意向对话→成交率 8%：")
    add_table(doc, ["测算项", "数值", "说明"], [
        ["月意向对话", "≈ 2,520", "来自 60 台配置产能估算"],
        ["月成交客户", "≈ 202", "2,520 × 8% 成交率"],
        ["月毛利", "≈ $30,240", "201 × 300 × 50%"],
        ["月运营成本", "≈ $4,000", "月服务费 3,000 + 代理/SIM/电费估算"],
        ["月净利估算", "≈ $26,240", "月毛利 − 运营成本"],
        ["软件部署费", f"{DEPLOY} {D.FEES['currency']}", "一次性；硬件按当日市场价自行采购、未计入"],
        ["部署费回本", "≈ 0.5 个月", "软件部署费 ÷ 月净利"],
    ], widths=[3.5, 3.5, 9], right_cols=[])
    add_callout(doc, "说明",
                "以上为理论测算示例，实际结果取决于市场、行业、客单价、话术质量与运营投入，不构成收益承诺。"
                "官网方案页提供在线交互测算器，可自行调整参数实时估算。", L_CYAN)

    # ── 十、实施交付计划 ──
    add_heading(doc, "十、实施与交付计划", 1)
    add_table(doc, ["阶段", "时间", "交付内容"], [
        ["签约与需求确认", "Day 0", "确认目标市场、平台、语言、话术方向；硬件到位或代采下单"],
        ["环境与集群搭建", "Day 1–3", "主控/Worker 部署，设备接入纳管，网络/代理/VPN 配置，投屏群控调通"],
        ["业务配置与联调", "Day 4–7", "引流链路、AI 话术、知识库、反风控参数配置；端到端联调全绿"],
        ["养号与灰度上线", "Day 7–10", "新号冷启动养号，小批量灰度跑通「触达→承接→引流」闭环"],
        ["放量与运营优化", "持续", "数据复盘、A/B 实验、策略调优，逐步扩大投产规模"],
    ], widths=[3.5, 2.5, 10])

    # ── 十一、服务与 SLA ──
    add_heading(doc, "十一、服务与 SLA", 1)
    add_para(doc, "月服务费覆盖以下持续保障内容：")
    add_table(doc, ["保障项", "说明"], [
        ["关键故障响应", "系统级故障 / 后台不可用，优先响应处置，快速恢复"],
        ["平台适配跟进", "目标平台 UI 更新导致的选择器失效，跟进修复适配"],
        ["反风控调优", "根据账号存活数据持续优化行为参数与配额策略"],
        ["版本迭代", "系统版本更新与现有功能范围内的小幅优化"],
        ["运营支持", "使用答疑、话术 / 策略建议、数据解读支持"],
    ], widths=[4, 12])

    # ── 十二、风险与合规 ──
    add_heading(doc, "十二、风险与合规声明", 1)
    add_heading(doc, "客户责任边界", 3)
    for t in [
        "账号需合法获取，操作须符合目标地区法律与平台条款。",
        "AI 生成内容建议人工最终把关，全程可在后台追溯。",
        "因违反平台条款导致的账号风控风险由客户自行承担。",
    ]:
        add_bullet(doc, t)
    add_heading(doc, "我方保障", 3)
    for t in [
        "系统私有化部署，数据完全隔离，我方不留存客户业务数据。",
        "内置反风控与合规护栏，最大化账号存活与稳定性。",
        "本方案面向正当出海营销获客，不支持任何诈骗 / 违法用途。",
    ]:
        add_bullet(doc, t)

    # ── 十三、常见问题 ──
    add_heading(doc, "十三、常见问题（FAQ）", 1)
    faqs = [
        (f"部署费 {DEPLOY} {D.FEES['currency']} 具体包含哪些？",
         "包含现有成熟功能范围内的完整系统私有化部署、TikTok/Facebook/Messenger 引流链路配置、"
         "设备集群接入与投屏群控、AI 客服话术与多语言知识库初始化、反风控与代理配置、"
         "数据看板与告警接入、端到端联调，以及 1 对 1 运营培训。一次性收取。"),
        (f"月服务费 {MONTH} {D.FEES['currency']} 是必须的吗？",
         "建议保留。社交平台 UI 和风控策略持续变化，月服务费用于保障系统稳定运行、平台适配跟进、"
         "反风控调优、版本迭代与优先技术支持，是自动化系统长期稳定投产的关键保障。"),
        ("硬件是你们提供还是我们自己买？",
         "硬件属客户一次性资产投入，独立于软件费用。可按配置清单自行采购，也可委托我方代采整备。"
         "由于硬件价格每日波动，本方案只给出配置与数量、不标注价格，实际以采购当日市场价为准。"),
        ("除了 TikTok 和 Facebook，还能做别的平台吗？",
         "可以。系统采用统一的多平台动作引擎，Instagram / WhatsApp / Telegram / X / LinkedIn / LINE 等"
         "均已具备同栈基础能力。若需作为正式引流渠道投产，按需求二次开发接入，单独评估报价。"),
        ("一台电脑能带多少部手机？",
         "每台工作机通过独立供电的工业 USB HUB 可稳定管理约 15–16 台真机；规模更大时通过增加 Worker "
         "节点横向扩展，单主控集群可管理 200+ 台设备。"),
        ("账号会不会被封？",
         "系统内置完整反风控栈（拟人行为、配额限流、冷却、代理伪装、指纹隔离、渐进放量等）最大化账号存活率。"
         "但社交平台风控客观存在，任何自动化都无法 100% 保证不被风控，建议合规操作 + 账号分层运营。"),
        ("多久能上线出效果？",
         "标准实施约 7–10 天完成部署、配置与灰度上线；新账号约 7 天养号期后进入正常触达节奏，"
         "通常 2–4 周可看到稳定引流与转化数据，随后进入放量优化阶段。"),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run("Q：" + q)
        _set_font(r, size=10.5, bold=True, color=INDIGO)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run("A：" + a)
        _set_font(r2, size=10)

    # ── 联系 ──
    add_heading(doc, "联系我们", 1)
    add_para(doc, "无界科技 BOUNDLESS —— 让沟通，无界。")
    add_para(doc, "Telegram 增长顾问：@ai_zkw")
    add_para(doc, "官网：https://usdt2026.cc")
    add_para(doc, "在线交互方案页：https://usdt2026.cc/proposal/")

    # ── 页脚 ──
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("无界科技 BOUNDLESS · 全球全自动引流系统部署方案书 · 商业机密 Confidential")
        _set_font(r, size=8, color="999999")

    _save(doc)


def _save(doc):
    """原子写入：先写临时文件，成功后 os.replace 覆盖正式文件（避免半截文件）。
    正式文件被占用（多为 Word 打开）时，回退到固定的 *_最新.docx 并提示；
    正式文件可写时，顺手清理历史遗留的 *_最新.docx 副本，避免重名混淆。"""
    alt = OUT.with_name(OUT.stem + "_最新" + OUT.suffix)
    fd, tmp = tempfile.mkstemp(suffix=".docx", dir=str(OUT.parent))
    os.close(fd)
    doc.save(tmp)
    try:
        os.replace(tmp, OUT)              # 原子覆盖
        print("SAVED:", OUT)
        if alt.exists():                  # 清理旧的回退副本
            try:
                alt.unlink()
                print("CLEANED duplicate:", alt.name)
            except OSError:
                pass
    except PermissionError:
        os.replace(tmp, alt)
        print("SAVED (正式文件被占用，已另存):", alt)
        print("  → 关闭 Word 中打开的旧文件后重跑本脚本，即可覆盖回正式文件名并自动清理副本。")


if __name__ == "__main__":
    build()
