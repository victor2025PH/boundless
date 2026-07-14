from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r"D:\workspace\欧美市场获客与AI转化系统改造评估报告.docx")

TITLE_COLOR = "1F4E79"
ACCENT = "D9EAF7"
DARK = "17365D"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
LIGHT_RED = "FCE4D6"
GRAY = "F2F2F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_fill=TITLE_COLOR):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], header_fill)
        if widths:
            hdr[i].width = Cm(widths[i])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_idx % 2 == 1:
                set_cell_shading(cells[i], "F8FBFD")
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(TITLE_COLOR if level <= 2 else DARK)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def add_callout(doc, title, body, fill=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(DARK)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)
    doc.add_paragraph()


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_doc():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run("欧美市场获客与 AI 转化系统\n改造评估报告")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于 mobile-auto0423 与 telegram-mtproto-ai 项目现状")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("666666")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(180)
    r = meta.add_run("文档用途：老板/合伙人内部决策\n版本：V1.0\n日期：2026-04-30")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "执行摘要", 1)
    add_callout(
        doc,
        "核心判断",
        "现有两个项目已经具备从社交平台触达到 AI 客服承接转化的基础能力。建议以 MVP 方式先验证欧美市场 CAC、线索质量和账号稳定性，再决定是否规模化投入。",
        LIGHT_GREEN,
    )
    add_table(
        doc,
        ["维度", "判断"],
        [
            ["技术可行性", "可行，已有手机自动化、AI 回复、漏斗、handoff、监控等基础能力"],
            ["商业建议", "先做 MVP，避免在未验证 CAC 前重投入"],
            ["推荐市场", "优先 US/UK/CA/AU 英语市场，暂不建议一开始做多语言欧洲"],
            ["推荐平台", "Facebook/Instagram Lead + Messenger 承接优先，TikTok 用于素材测试和低成本流量"],
            ["MVP 周期", "3-5 周"],
            ["可投放运营版周期", "6-10 周"],
            ["规模化周期", "10-16 周"],
            ["MVP 开发预算", "USD 8,000-20,000"],
            ["冷启动广告测试预算", "USD 2,000-8,000"],
            ["最大风险", "平台政策、欧美隐私合规、账号稳定性、广告线索质量"],
        ],
        widths=[4, 12],
    )

    add_heading(doc, "一、现有项目能力盘点", 1)
    add_heading(doc, "1. mobile-auto0423：获客执行层", 2)
    add_para(doc, "该项目定位为真机移动端社交平台自动化系统，基于 ADB 与 uiautomator2 控制 Android 设备，具备多设备集群、任务调度、设备健康检查、Facebook/Messenger/TikTok 自动化、漏斗统计、A/B 实验和风控检测等能力。")
    add_table(doc, ["能力", "当前价值", "对欧美改造的意义"], [
        ["Facebook 加好友/打招呼", "可主动触达目标客户", "可改造为欧美目标人群触达链路"],
        ["Messenger 收件箱监控与 AI 回复", "可承接私信沟通", "可承接广告私信与自然回复"],
        ["TikTok 内容/互动相关模块", "已有平台入口", "可作为低成本素材测试与流量入口"],
        ["漏斗与 ROI 面板", "可观察转化数据", "可扩展到广告 CAC/CPL/ROAS"],
        ["A/B 实验", "可测试话术与策略", "可用于欧美话术、素材、落地页实验"],
        ["设备集群与风控", "支撑多账号运行", "决定规模化稳定性与成本"],
    ], widths=[4.5, 5.5, 6])

    add_heading(doc, "2. telegram-mtproto-ai：AI 客服与转化承接层", 2)
    add_para(doc, "该项目定位为多平台 AI 客服主骨架，包含 Telegram、LINE、Messenger RPA runner，contacts/handoff 跨平台客户系统，知识库回复、Web 后台、监控与审计能力。")
    add_table(doc, ["能力", "当前价值", "对欧美改造的意义"], [
        ["AI 回复栈", "支持自动客服与知识库回复", "可建设英文销售客服与 FAQ"],
        ["contacts/handoff", "统一客户、旅程、转人工", "可承接广告线索到真人销售"],
        ["多平台 runner", "支持 Telegram/LINE/Messenger", "可按欧美客户偏好切换承接渠道"],
        ["Web 后台", "可查看 funnel、handoff、monitoring", "可扩展为运营决策后台"],
        ["observability", "指标、审计、告警", "支撑规模化运行与风险控制"],
    ], widths=[4.5, 5.5, 6])

    add_heading(doc, "二、欧美市场改造需求", 1)
    add_heading(doc, "1. 市场与人群本地化", 2)
    add_para(doc, "欧美市场不能简单翻译现有话术，需要重新设计 persona、渠道偏好、信任建立方式和销售节奏。建议优先从英语国家开始。")
    add_table(doc, ["市场", "优先级", "理由"], [
        ["美国", "高", "市场大、广告生态成熟、数据样本充足"],
        ["英国", "中高", "英语市场，合规要求较高但可复用美国话术"],
        ["加拿大", "中", "英语用户多，广告竞争相对可控"],
        ["澳大利亚", "中", "英语市场，客单价通常较好"],
        ["欧洲非英语国家", "后置", "需要多语言、本地合规和客服资源"],
    ], widths=[4, 3, 9])

    add_heading(doc, "2. 英文话术与销售流程", 2)
    add_table(doc, ["模块", "内容"], [
        ["Opening", "英文开场白、广告私信欢迎语、自然对话开场"],
        ["Qualification", "预算、需求、时间、地区、痛点识别"],
        ["Objection Handling", "价格、信任、隐私、效果、售后等反异议"],
        ["CTA", "预约、领取资料、加入 Telegram/WhatsApp、试用、购买"],
        ["Handoff", "高意向转人工、低意向继续培育"],
        ["Follow-up", "24h/72h/7d 跟进节奏"],
    ], widths=[5, 11])

    add_heading(doc, "3. 合规与隐私", 2)
    add_callout(doc, "最低要求", "欧美市场必须把合规作为产品能力，而不是后补文案。隐私政策、用户同意、退订、频控和数据留存策略应在 MVP 阶段就具备。", LIGHT_YELLOW)
    add_table(doc, ["项目", "要求"], [
        ["隐私政策", "落地页和表单必须可访问 Privacy Policy"],
        ["Consent", "表单或私信中明确用户同意被联系"],
        ["Opt-out", "支持 STOP/Unsubscribe/Do not contact"],
        ["数据留存", "规定线索数据保存周期与删除机制"],
        ["数据导出/删除", "至少在后台保留人工处理流程"],
        ["联系频控", "同一客户 24h/7d 最大触达次数"],
    ], widths=[5, 11])

    add_heading(doc, "三、开发方案与里程碑", 1)
    add_heading(doc, "阶段 1：MVP 验证版，3-5 周", 2)
    add_table(doc, ["模块", "交付物", "预估周期"], [
        ["英文 persona 与话术", "US/UK 英文开场、FAQ、转化话术", "3-5 天"],
        ["基础合规", "opt-out、频控、隐私提示、数据字段", "3-5 天"],
        ["广告来源字段", "UTM/source/campaign/ad 字段入库", "3-5 天"],
        ["漏斗报表", "lead → chat → qualified → handoff → converted", "5-7 天"],
        ["AI 客服配置", "英文知识库、反异议、handoff 规则", "5-7 天"],
        ["真机链路测试", "Facebook/Messenger/TikTok 关键路径测试", "5-10 天"],
    ], widths=[4.5, 8, 3])
    add_heading(doc, "MVP 验收标准", 3)
    add_table(doc, ["指标", "目标"], [
        ["线索可归因率", "≥ 90%"],
        ["AI 首响时间", "≤ 60 秒，视平台限制可放宽"],
        ["opt-out 生效率", "100%"],
        ["漏斗数据完整率", "≥ 90%"],
        ["跑通样本量", "至少 100-300 条线索"],
    ], widths=[6, 10])

    add_heading(doc, "阶段 2：可投放运营版，6-10 周", 2)
    add_table(doc, ["模块", "交付物"], [
        ["Pixel/CAPI/TikTok Pixel", "广告事件回传与优化信号"],
        ["Campaign 看板", "国家、平台、素材、渠道维度统计"],
        ["A/B 实验", "话术、CTA、落地页、客服策略实验"],
        ["客服工作台", "高意向线索、待跟进、超时提醒"],
        ["风控策略", "账号、设备、频率、敏感词、冷却策略"],
        ["成本追踪", "LLM 成本、设备成本、代理成本、广告成本"],
    ], widths=[6, 10])

    add_heading(doc, "阶段 3：规模化版本，10-16 周", 2)
    add_table(doc, ["模块", "交付物"], [
        ["多国家 domain pack", "US/UK/CA/AU 与欧洲语言包"],
        ["统一客户画像", "跨平台客户 ID、旅程、标签与成交历史"],
        ["自动预算建议", "按 CAC/LTV 推荐增减预算"],
        ["多账号与设备调度", "动态配额、异常隔离、任务优先级"],
        ["合规审计", "联系记录、退订记录、数据删除记录"],
        ["运营周报", "自动生成 campaign 复盘和 ROI 报表"],
    ], widths=[6, 10])

    add_heading(doc, "四、开发成本估算", 1)
    add_table(doc, ["版本", "周期", "开发预算", "适合情况"], [
        ["MVP", "3-5 周", "USD 8,000-20,000", "验证欧美获客是否跑得通"],
        ["可投放运营版", "6-10 周", "USD 20,000-50,000", "准备持续投放和运营复盘"],
        ["规模化版", "10-16 周", "USD 50,000-120,000+", "多国家、多平台、多账号放量"],
    ], widths=[3.5, 3, 4.5, 5])
    add_heading(doc, "非开发成本", 2)
    add_table(doc, ["成本项", "月度参考"], [
        ["广告测试预算", "USD 2,000-8,000 起"],
        ["Android 真机/云手机", "USD 200-2,000+"],
        ["账号/SIM/验证资源", "USD 200-2,000+"],
        ["代理/VPN/IP 环境", "USD 300-3,000+"],
        ["LLM API 成本", "USD 100-2,000+"],
        ["云服务/监控/存储", "USD 100-1,000+"],
        ["人工客服/销售", "视团队成本而定"],
    ], widths=[7, 9])

    add_heading(doc, "五、Facebook 与 TikTok 获客成本分析", 1)
    add_heading(doc, "Facebook/Instagram Ads", 2)
    add_table(doc, ["指标", "常见区间", "说明"], [
        ["CPC", "USD 0.5-3.0", "普通消费类较常见"],
        ["高竞争 CPC", "USD 3-10+", "金融、教育、B2B、法律等行业"],
        ["CPL", "USD 5-30", "Lead Form 或落地页表单常见区间"],
        ["高质量 CPL", "USD 30-150+", "B2B、高客单价、强筛选人群"],
        ["私信线索", "USD 3-25", "取决于素材和国家"],
    ], widths=[4, 4, 8])
    add_heading(doc, "TikTok Ads", 2)
    add_table(doc, ["指标", "常见区间", "说明"], [
        ["CPC", "USD 0.2-1.5", "通常低于 Facebook"],
        ["CPL", "USD 3-20", "前期可能便宜但质量波动较大"],
        ["高质量 CPL", "USD 20-80+", "需要更强素材筛选和客服承接"],
        ["素材迭代频率", "高", "TikTok 素材衰减快"],
    ], widths=[4, 4, 8])
    add_heading(doc, "冷启动预算建议", 2)
    add_table(doc, ["测试阶段", "预算", "目标"], [
        ["第 1 周", "USD 1,000-2,000", "测试 3-5 个素材方向和基础 CPL"],
        ["第 2 周", "USD 1,000-3,000", "保留胜出素材，验证线索质量"],
        ["第 3-4 周", "USD 2,000-5,000", "小规模放量，观察 CAC 与成交"],
        ["合计", "USD 2,000-8,000", "获取是否值得继续投入的初步判断"],
    ], widths=[4, 4, 8])

    add_heading(doc, "六、ROI 测算模型", 1)
    add_callout(doc, "基础公式", "CPL = 广告花费 / 线索数；CAC = 广告花费 / 成交客户数；ROAS = 成交收入 / 广告花费。若客单价低，必须压低 CPL 并提升客服转化率；若客单价高，可接受更高 CPL。", ACCENT)
    add_para(doc, "以下示例假设：月广告预算 USD 5,000，客单价 USD 300，毛利率 60%。")
    add_table(doc, ["场景", "CPL", "线索数", "有效率", "有效线索", "成交率", "成交数", "CAC", "毛利", "结果"], [
        ["保守", "30", "167", "30%", "50", "5%", "3", "1,667", "540", "不建议放量"],
        ["基准", "15", "333", "40%", "133", "8%", "11", "455", "1,980", "需优化客单或转化"],
        ["乐观", "8", "625", "50%", "313", "12%", "38", "132", "6,840", "可考虑放量"],
    ], widths=[2, 1.6, 1.7, 1.8, 1.8, 1.8, 1.6, 1.8, 1.8, 3])

    add_heading(doc, "七、风险与控制策略", 1)
    add_table(doc, ["风险", "影响", "控制策略"], [
        ["平台政策风险", "账号限制、广告拒登、触达下降", "保守频控、人工审核、合规话术、账号分层"],
        ["欧美隐私合规", "投诉、法律风险、广告账户受限", "隐私政策、opt-out、数据留存、删除流程"],
        ["账号与设备稳定性", "业务中断、成本上升", "健康检查、设备池隔离、异常自动降级"],
        ["广告素材质量", "CPL 高、线索差", "每周素材迭代、A/B 测试、素材疲劳监控"],
        ["AI 回复失误", "客户流失、投诉", "高意向转人工、敏感问题不自动回复、审计日志"],
        ["归因不准确", "无法判断 ROI", "UTM、Pixel、事件回传、跨平台客户 ID"],
    ], widths=[4, 5, 7])

    add_heading(doc, "八、推荐决策", 1)
    add_callout(doc, "是否建议立项", "建议立项，但建议以 MVP 方式立项，而不是直接做大规模版本。最大不确定性不是技术，而是 CAC、线索质量和平台稳定性，必须用小预算实测。", LIGHT_GREEN)
    add_heading(doc, "建议首期预算", 2)
    add_table(doc, ["项目", "建议预算"], [
        ["MVP 开发", "USD 8,000-20,000"],
        ["广告测试", "USD 2,000-8,000"],
        ["设备/账号/代理/云服务", "USD 1,000-5,000"],
        ["合计", "USD 11,000-33,000"],
    ], widths=[8, 8])
    add_heading(doc, "30/60/90 天行动计划", 2)
    add_table(doc, ["时间", "目标", "关键动作"], [
        ["0-30 天", "MVP 跑通", "英文话术、基础合规、广告来源字段、基础漏斗、首批投放"],
        ["31-60 天", "数据验证", "素材 A/B、客服转化优化、Pixel/CAPI、有效线索定义"],
        ["61-90 天", "小规模放量", "多 campaign、客服工作台、预算调整、账号/设备扩容"],
    ], widths=[3.5, 4, 8.5])

    add_heading(doc, "九、决策前需要确认的问题", 1)
    for item in [
        "具体业务类型是什么：B2C、B2B、订阅、电商、服务预约，还是其他？",
        "客单价、毛利率和客户生命周期价值大概是多少？",
        "首选国家是美国，还是其他英语国家？",
        "转化目标是留资、加 Telegram/WhatsApp、预约、购买，还是人工销售跟进？",
        "是否已有广告账户、素材、落地页、隐私政策和客服团队？",
        "是否接受前期用 USD 2,000-8,000 广告预算测试真实 CAC？",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "附录：功能模块映射", 1)
    add_table(doc, ["业务环节", "主要项目", "现有基础", "改造重点"], [
        ["目标触达", "mobile-auto0423", "Facebook/TikTok/Messenger 自动化", "欧美频控、目标人群、英文话术"],
        ["私信承接", "mobile-auto0423 + telegram-mtproto-ai", "AI 回复、Messenger runner", "英文客服、反异议、转人工"],
        ["客户画像", "telegram-mtproto-ai", "contacts/journey", "source/campaign/国家/意向标签"],
        ["引流转化", "telegram-mtproto-ai", "handoff、LINE/TG", "WhatsApp/Telegram/预约链路"],
        ["数据看板", "两项目", "funnel、analytics、monitoring", "CAC/CPL/ROAS、Campaign 维度"],
        ["风险控制", "两项目", "rate limit、risk、health monitor", "GDPR/CCPA、opt-out、平台政策"],
    ], widths=[3.5, 4.5, 4, 4])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("欧美市场获客与 AI 转化系统改造评估报告 | 内部决策版")
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("888888")

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(str(OUT))
