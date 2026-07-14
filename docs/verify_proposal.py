# -*- coding: utf-8 -*-
"""
一致性校验：确保「网页 index.html」与「Word 方案书」与单一事实源 proposal_data 保持一致，
并确认已移除的硬件价格没有回流。CI/改稿后手动跑一次即可，避免两处数字漂移。

用法：  python verify_proposal.py
退出码：0 = 全部通过；1 = 有不一致项。
"""
import sys
from pathlib import Path

try:                       # Windows 控制台默认 GBK，强制 UTF-8 避免符号报错
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
import proposal_data as D  # noqa: E402

HTML = Path(r"D:\workspace\ai-p0-integration\website\public\proposal\index.html")
DOCX = HERE / "无界全球全自动引流系统_部署方案书.docx"

checks = []          # (name, ok, detail)


def add(name, ok, detail=""):
    checks.append((name, bool(ok), detail))


def load_docx_text(path: Path) -> str:
    from docx import Document
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            parts += [c.text for c in r.cells]
    return "\n".join(parts)


def run():
    html = HTML.read_text(encoding="utf-8") if HTML.exists() else ""
    docx = load_docx_text(DOCX) if DOCX.exists() else ""
    add("网页文件存在", bool(html), str(HTML))
    add("Word 文件存在", bool(docx), str(DOCX))
    if not html or not docx:
        return report()

    # 1) 费用一致（两处都应含部署费、月服务费数字）
    for label, val in [("部署费", D.FEES["deploy"]), ("月服务费", D.FEES["month"])]:
        add(f"网页含{label} {val}", val in html)
        add(f"Word 含{label} {val}", val in docx)

    # 2) 三档真机数量一致
    for tier in D.TIERS:
        n = str(tier["phones"])
        add(f"网页含{tier['label']} {n}台", n in html)
        add(f"Word 含{tier['label']} {n}台", n in docx)

    # 3) 九大平台名称都在
    for name, _cap, _st in D.PLATFORMS:
        add(f"网页含平台 {name}", name in html)
        add(f"Word 含平台 {name}", name in docx)

    # 4) 平台状态计数（3 已上线 / 6 可扩展）
    live = sum(1 for *_, st in D.PLATFORMS if st == "live")
    dev = sum(1 for *_, st in D.PLATFORMS if st == "dev")
    add(f"网页“已上线”≥{live}", html.count("已上线") >= live, f"实测 {html.count('已上线')}")
    add(f"网页“可扩展”≥{dev}", html.count("可扩展") >= dev, f"实测 {html.count('可扩展')}")
    add(f"Word “已上线”≥{live}", docx.count("已上线") >= live)
    add(f"Word “可扩展”≥{dev}", docx.count("可扩展") >= dev)

    # 5) 产能默认值（网页默认展示 60 台的月意向对话）
    p60 = D.productivity(60)
    add(f"网页含 60 台月意向对话 {D.fmt(p60['leads'])}", D.fmt(p60["leads"]) in html)

    # 6) 已移除的硬件价格 token 不应回流
    for tok in D.FORBIDDEN_PRICE_TOKENS:
        add(f"网页无价格残留「{tok}」", tok not in html)
        add(f"Word 无价格残留「{tok}」", tok not in docx)

    return report()


def report():
    passed = sum(1 for _, ok, _ in checks if ok)
    total = len(checks)
    print("=" * 56)
    print(f" 方案一致性校验：{passed}/{total} 通过")
    print("=" * 56)
    for name, ok, detail in checks:
        mark = "✓" if ok else "✗"
        line = f" [{mark}] {name}"
        if detail and not ok:
            line += f"  ({detail})"
        print(line)
    fails = [n for n, ok, _ in checks if not ok]
    if fails:
        print("\n不通过项：", "; ".join(fails))
        return 1
    print("\n全部通过：网页 / Word / 数据源三方一致。")
    return 0


if __name__ == "__main__":
    sys.exit(run())
