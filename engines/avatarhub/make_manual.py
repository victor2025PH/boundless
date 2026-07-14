# -*- coding: utf-8 -*-
"""生成《AvatarHub 功能使用说明》Word 文档（精美排版）。
运行：<facefusion>\python.exe make_manual.py  → 产出 AvatarHub_功能使用说明.docx
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 品牌配色 ──
BLUE   = RGBColor(0x25, 0x63, 0xEB)   # 主蓝
DARK   = RGBColor(0x1E, 0x29, 0x3B)   # 深色标题
GREY   = RGBColor(0x6B, 0x72, 0x80)   # 次要灰
GREEN  = RGBColor(0x15, 0x80, 0x3D)
AMBER  = RGBColor(0xB4, 0x53, 0x09)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FONT   = "微软雅黑"
FONT_T = "微软雅黑"


def set_run(run, name=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(a), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = color


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def para(doc, text, size=10.5, color=None, bold=False, align=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color or DARK)
    return p


def bullet(doc, text, lead="", size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if lead:
        r = p.add_run(lead); set_run(r, size=size, bold=True, color=BLUE)
    r2 = p.add_run(text); set_run(r2, size=size, color=DARK)
    return p


def h1(doc, text):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); set_run(r, name=FONT_T, size=17, bold=True, color=BLUE)
    # 底部色条
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '2563EB')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_run(r, name=FONT_T, size=13, bold=True, color=DARK)
    return p


def code(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run(text); set_run(r, name="Consolas", size=9.5, color=RGBColor(0x0B,0x3D,0x91))
    # 浅灰底
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F1F5F9')
    pPr.append(shd)
    return p


def table(doc, headers, rows, widths=None, header_fill='2563EB'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_text(c, hd, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(c, header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], str(val), size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


def callout(doc, title, text, fill='EFF6FF', bar='2563EB', tcolor=BLUE):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; shade(cell, fill)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  "); set_run(r, size=10, bold=True, color=tcolor)
    r2 = p.add_run(text); set_run(r2, size=10, color=DARK)
    # 左色条
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders'); left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'24'); left.set(qn('w:space'),'0'); left.set(qn('w:color'),bar)
    borders.append(left); tcPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    t = OxmlElement('w:t'); t.text = '（在 Word 中按 F9 或右键“更新域”即可生成目录）'
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    for e in (f1, instr, f2, t, f3): run._r.append(e)


# ════════════════════════════════════════════════════════════
doc = Document()
# 默认正文样式
st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
for sec in doc.sections:
    sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

# ── 封面 ──
for _ in range(3): doc.add_paragraph()
band = doc.add_table(rows=1, cols=1); band.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = band.rows[0].cells[0]; shade(bc, '2563EB')
bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(18); bp.paragraph_format.space_after = Pt(18)
r = bp.add_run("AvatarHub"); set_run(r, size=40, bold=True, color=WHITE)
bp2 = bc.add_paragraph(); bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp2.paragraph_format.space_after = Pt(16)
r = bp2.add_run("实时数字人系统 · 功能使用说明"); set_run(r, size=16, bold=True, color=WHITE)

doc.add_paragraph()
para(doc, "语音克隆 · 活体口型 · 多语种对话 · WebRTC / OBS 直播",
     size=13, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "说话 → 语音识别 → 大模型 → 克隆音合成 → 流式活体口型 → 虚拟摄像头",
     size=10.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(6): doc.add_paragraph()
info = doc.add_table(rows=3, cols=2); info.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [("版本", "v1.0"), ("文档日期", datetime.date.today().strftime("%Y-%m-%d")),
        ("适用对象", "创作者 / 主播 / 集成商 / 运维")]
for i,(k,v) in enumerate(data):
    cell_text(info.rows[i].cells[0], k, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(info.rows[i].cells[0], '1E293B')
    cell_text(info.rows[i].cells[1], v, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for r_ in info.rows:
    r_.cells[0].width = Inches(1.4); r_.cells[1].width = Inches(3.2)
doc.add_page_break()

# ── 目录 ──
h1(doc, "目录")
add_toc(doc)
doc.add_page_break()

# ── 1. 产品概述 ──
h1(doc, "一、产品概述")
para(doc, "AvatarHub 是一套运行在本机 / 私有局域网的实时数字人平台：上传一张照片与一段参考语音，"
          "即可生成“会说话、有表情、能实时对话、可直播”的数字人。全链路本地化部署，数据不出机，"
          "并内置合规溯源（水印 + 数字签名）。")
h2(doc, "核心能力")
bullet(doc, "零样本声音克隆（Fish-Speech），几秒参考音即可复刻音色；", "• ")
bullet(doc, "活体数字人（LivePortrait + MuseTalk）：会摆头、眨眼、有表情，非“死图对口型”；", "• ")
bullet(doc, "实时对话：语音/文字输入，低延迟流式应答，可中途“抢话”打断；", "• ")
bullet(doc, "多语种、可直播：一键推流到 OBS 虚拟摄像头 / WebRTC；", "• ")
bullet(doc, "实时高清（HD）口型：单卡即可 25fps 高清直播；", "• ")
bullet(doc, "观众互动：观众提问→排队→数字人串行应答，公屏 + 精彩问答回放。", "• ")

h2(doc, "系统架构（中枢 + 微服务）")
para(doc, "中央编排器 AvatarHub（端口 9000）聚合下列微服务，各自运行在隔离环境中：", space_after=4)
table(doc,
      ["服务", "作用", "核心链路"],
      [["AvatarHub 中枢", "角色管理 / 编排 / API / 控制台", "✅"],
       ["克隆音 TTS", "Fish-Speech 零样本克隆", "✅"],
       ["语音转文字", "Whisper 语音识别", "✅"],
       ["口型同步", "MuseTalk + LivePortrait 活体", "✅"],
       ["广播中枢", "OBS 虚拟摄像头 + WebRTC", "✅"],
       ["换脸 / 情感TTS / 唱歌 / 发型 等", "可选扩展（按需启动）", "⬜"]],
      widths=[1.8, 3.4, 1.0])
callout(doc, "提示", "标 ✅ 为实时对话最小集，开机即起；标 ⬜ 为可选扩展，需要时设 START_EXTRAS=1 启动，"
                    "默认不启动属正常，不影响对话与直播。")

# ── 2. 快速开始 ──
h1(doc, "二、快速开始")
h2(doc, "启动")
bullet(doc, "桌面启动器（推荐）：双击 launcher.bat，图形界面一键“启动 / 停止 / 重启 / 体检”，并实时显示各服务就绪状态。", "1. ")
bullet(doc, "命令行：运行 start_all_services.bat 一键拉起核心链路（就绪后自动放行）。", "2. ")
h2(doc, "进入后台控制台")
para(doc, "浏览器访问统一控制台：", space_after=2); code(doc, "http://127.0.0.1:9000/ui")
para(doc, "手机端对话页（同一局域网）：", space_after=2); code(doc, "http://<本机局域网IP>:9000/phone")
callout(doc, "首次激活", "首次激活一个新角色需约 10–20 秒做活体/口型预计算，控制台会显示“数字人准备中…”横幅，"
                      "完成后自动转“已就绪”，期间首句走静态回退。同一角色之后激活仅需约 0.4 秒。", fill='FEF3C7', bar='B45309', tcolor=AMBER)

# ── 3. 控制台功能详解 ──
h1(doc, "三、控制台功能详解")
para(doc, "控制台按“角色 / 创作 / 运营”三组组织，左侧切换。各功能说明如下：", space_after=6)

feats = [
    ("🎭 角色库", "创建 / 导入 / 激活数字人角色。每个角色含照片、克隆音、人设（system_prompt）、效果参数。"
                "激活后即作为当前对话与直播的形象与音色。支持音质评分（cosine）、黄金出厂包一键恢复。"),
    ("🧬 克隆", "上传参考语音克隆音色；可多段参考融合提升稳定性。提供录音质量自检（信噪比/电平/时长）与录制建议。"),
    ("🎙️ 语音", "文本→语音试听与对比：比较当前角色两种 TTS 引擎的音色与自然度；调试克隆效果与情感语音。"),
    ("🎵 唱歌", "歌声合成（可选扩展）：用克隆音色演唱。需启动 singing 服务。"),
    ("📦 批量", "批量文本转语音/口型：每行一句，支持多角色语法 [角色名] 文本；长任务带进度条。"),
    ("📊 看板", "运行总览：服务健康、性能（GPU/显存/延迟）、对话指标、容量与排队状态实时可视。"),
    ("📡 开播", "一键推流：把数字人画面输出到 OBS 虚拟摄像头 / WebRTC；管理摄像头、变声设备、实时状态。"),
    ("📜 历史", "对话与生成历史：检索、收藏、删除；统计与增量加载。"),
    ("✅ 交付体检", "一键自检面板：现跑联机体检（doctor）并展示最近一次交付自检结论（红/黄/绿），逐项列出服务/角色/容量/"
                 "流式就绪状态。可选扩展未启用属正常、不计入结论。"),
    ("📋 日志", "系统日志实时查看，按 INFO/WARNING/ERROR 过滤，可自动刷新。"),
    ("⚙️ 设置", "运行参数、指标、知识库（RAG）等配置入口。"),
]
for name, desc in feats:
    h2(doc, name)
    para(doc, desc, space_after=4)

# ── 4. 实时对话与手机端 ──
h1(doc, "四、实时对话与手机端")
bullet(doc, "网页对话：控制台内直接输入文本/语音与数字人对话，支持流式应答（边说边出）。", "• ")
bullet(doc, "手机端 /phone：扫码或输入局域网地址，手机即可语音对话；免按键（VAD）模式下“开口即打断”。", "• ")
bullet(doc, "开场白缓存：常用开场即时应答，进一步压低首句延迟。", "• ")
callout(doc, "实测延迟", "单卡（RTX 5090）单路对话：文本→首帧 标清约 1.3 秒、高清约 4.5 秒；交互接近自然对话。")

# ── 5. 观众互动与直播 ──
h1(doc, "五、观众互动与直播")
bullet(doc, "观众提问 /ask：观众提交问题，进入公平队列，数字人逐条串行应答（可设自动应答 + 简答）。", "• ")
bullet(doc, "公屏 /wall：OBS 浏览器源叠加，展示“正在回答”的问题，支持人气（点赞）优先。", "• ")
bullet(doc, "精彩问答 /highlights：每条问答持久化，可导出竖版图卡（4:5 / 9:16 / 1:1）做短视频复用。", "• ")
para(doc, "开启方式：设置环境变量 AVATARHUB_AUDIENCE=1 后重启 Hub（默认关闭）。", space_after=6)

# ── 6. 容量与并发 ──
h1(doc, "六、容量与并发（重要）")
para(doc, "TTS 与口型共用 GPU，单卡无法真并发。系统采用“准入 + 公平队列”作为降级策略：", space_after=4)
table(doc, ["场景", "表现", "建议"],
      [["单卡 1 路对话", "首帧 标清~1.3s / 高清~4.5s，体验优秀", "默认形态"],
       ["单卡多路并发", "互相拖垮、延迟飙升、不可预期", "开准入排队"],
       ["真并发(多路独立)", "需多卡/多机", "SVC_LIPSYNC 多副本"]],
      widths=[1.8, 3.2, 1.6])
callout(doc, "推荐配置", "单卡设 CONV_MAX_CONCURRENT=auto + CONV_MAX_QUEUE=20：1 路实时数字人 + 观众排队串行应答，"
                       "超额路优雅排队并显示预计等待（ETA）。HD 直播机设 LIPSYNC_HD_PREWARM=1，首条高清句即满速。")

# ── 7. 商用部署：安全与授权 ──
h1(doc, "七、商用部署：安全与授权")
para(doc, "本机自用可全部留空；一旦暴露到局域网 / 手机 / 外网，按下表加固（详见 env_config.bat「COMMERCIAL DEPLOYMENT」段）。", space_after=4)
table(doc, ["项", "环境变量 / 命令", "作用"],
      [["管理面令牌", "AVATARHUB_API_TOKEN", "挡跨机改配置 / 操纵角色"],
       ["子服务令牌", "AVATARHUB_SERVICE_TOKEN / _ALLOW_IPS", "保护 GPU 算力口"],
       ["离线授权", "license_admin.py issue …", "按机器指纹绑定 + 有效期 + 档位"],
       ["强制授权", "AVATARHUB_LICENSE_ENFORCE=1", "无效授权只降级、绝不崩"]],
      widths=[1.3, 3.2, 2.0])
h2(doc, "授权激活三步")
code(doc, "1) 客户取指纹：  python license_admin.py fingerprint")
code(doc, "2) 厂商签发码：  python license_admin.py issue --machine <指纹> --edition pro --days 365")
code(doc, "3) 查看状态：    python license_admin.py status   或访问 /api/license/status")

# ── 8. 运维与自检 ──
h1(doc, "八、运维与自检")
bullet(doc, "就绪探针：python ready.py（退出码 0 = 核心链路全就绪）。", "• ")
bullet(doc, "联机体检：python doctor.py（服务 / 角色 / 容量 / 流式就绪）。", "• ")
bullet(doc, "一键交付自检：python deliver_check.py（0=可交付 / 1=可上线有警告 / 2=不可交付），结论也可在控制台“交付体检”tab 查看。", "• ")
bullet(doc, "守护与自愈：守护进程自动重启崩溃的核心服务；中枢具备多副本负载均衡与故障切换。", "• ")

# ── 9. 常见问题 ──
h1(doc, "九、常见问题（FAQ）")
faqs = [
    ("激活角色像“卡住”？", "首次激活新角色需 10–20 秒预计算，控制台有“数字人准备中”提示，属正常；同角色之后秒级就绪。"),
    ("交付体检里有 7 项“可选信息”？", "那是默认未启动的可选扩展服务（换脸/唱歌等），属预期，不影响对话与直播，也不计入交付结论。"),
    ("高清直播卡顿？", "确认设置 LIPSYNC_HD_PREWARM=1 做启动预热；单卡 HD 单路实时可达，多路需加卡。"),
    ("局域网别的机器能乱改我的角色？", "设置 AVATARHUB_API_TOKEN 管理面令牌即可封堵跨机写操作。"),
    ("断网还能对话吗？", "本地 LLM（Ollama）+ 本地 TTS/口型 全链路可离线；若默认用云端 LLM，建议保留本地兜底。"),
]
for q, a in faqs:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q：" + q); set_run(r, size=10.5, bold=True, color=BLUE)
    para(doc, "A：" + a, size=10.5, space_after=8)

# ── 页脚 ──
doc.add_paragraph()
foot = para(doc, "© AvatarHub · 实时数字人系统　本机 / 私有部署　数据不出机 · 合规溯源",
            size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

out = "AvatarHub_功能使用说明.docx"
doc.save(out)
print("已生成：" + out)
